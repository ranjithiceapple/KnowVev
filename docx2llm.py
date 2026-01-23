import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from pathlib import Path
import argparse
import sys
import json
import re
from collections import Counter
from datetime import datetime
import pandas as pd
try:
    import tiktoken
except ImportError:
    tiktoken = None


def convert_docx_to_markdown(docx_path, output_path=None):
    """
    Convert DOCX to Markdown format

    Args:
        docx_path: Path to the input DOCX file
        output_path: Path for the output markdown file (optional)

    Returns:
        markdown_text: The converted markdown text
    """
    docx_path = Path(docx_path)

    # Validate input file
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    if not docx_path.suffix.lower() in ['.docx', '.doc']:
        raise ValueError(f"File must be a DOCX: {docx_path}")

    print(f"Converting DOCX to Markdown: {docx_path.name}")
    print(f"{'='*70}")

    try:
        # Load document
        doc = docx.Document(docx_path)

        # Convert to markdown
        markdown_lines = []

        for element in iter_block_items(doc):
            if isinstance(element, Paragraph):
                md_line = paragraph_to_markdown(element)
                if md_line:
                    markdown_lines.append(md_line)

            elif isinstance(element, Table):
                md_table = table_to_markdown(element)
                if md_table:
                    markdown_lines.extend(md_table)
                    markdown_lines.append('')  # Add blank line after table

        markdown_text = '\n'.join(markdown_lines)

        # Determine output path
        if output_path is None:
            output_path = docx_path.with_suffix('.md')
        else:
            output_path = Path(output_path)

        # Save markdown file
        print(f"Saving markdown to: {output_path.name}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)

        # Calculate statistics
        lines = len(markdown_text.split('\n'))
        words = len(markdown_text.split())
        chars = len(markdown_text)

        print(f"{'='*70}")
        print("Conversion Statistics:")
        print(f"  Lines: {lines:,}")
        print(f"  Words: {words:,}")
        print(f"  Characters: {chars:,}")
        print(f"{'='*70}")
        print(f"SUCCESS: Markdown saved to {output_path}")

        return markdown_text

    except Exception as e:
        print(f"ERROR: Failed to convert DOCX")
        print(f"  {str(e)}")
        raise


def convert_docx_to_text(docx_path, output_path=None):
    """
    Convert DOCX to plain text format

    Args:
        docx_path: Path to the input DOCX file
        output_path: Path for the output text file (optional)

    Returns:
        text: The extracted plain text
    """
    docx_path = Path(docx_path)

    # Validate input file
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    print(f"Converting DOCX to Text: {docx_path.name}")
    print(f"{'='*70}")

    try:
        # Load document
        doc = docx.Document(docx_path)

        # Extract text
        text_lines = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)

        # Add table content
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_lines.append(row_text)
            text_lines.append('')  # Blank line after table

        text = '\n'.join(text_lines)

        # Determine output path
        if output_path is None:
            output_path = docx_path.with_suffix('.txt')
        else:
            output_path = Path(output_path)

        # Save text file
        print(f"Saving text to: {output_path.name}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

        # Statistics
        lines = len(text.split('\n'))
        words = len(text.split())

        print(f"{'='*70}")
        print("Extraction Statistics:")
        print(f"  Lines: {lines:,}")
        print(f"  Words: {words:,}")
        print(f"{'='*70}")
        print(f"SUCCESS: Text saved to {output_path}")

        return text

    except Exception as e:
        print(f"ERROR: Failed to convert DOCX")
        print(f"  {str(e)}")
        raise


def iter_block_items(parent):
    """
    Yield each paragraph and table child within parent, in document order
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_to_markdown(paragraph):
    """
    Convert a paragraph to markdown format
    """
    text = paragraph.text.strip()
    if not text:
        return ''

    # Check paragraph style
    style_name = paragraph.style.name if paragraph.style else ''

    # Handle headings
    if 'Heading 1' in style_name:
        return f"# {text}"
    elif 'Heading 2' in style_name:
        return f"## {text}"
    elif 'Heading 3' in style_name:
        return f"### {text}"
    elif 'Heading 4' in style_name:
        return f"#### {text}"
    elif 'Heading 5' in style_name:
        return f"##### {text}"
    elif 'Heading 6' in style_name:
        return f"###### {text}"

    # Handle lists
    elif 'List' in style_name:
        # Check if numbered or bullet
        if paragraph.style.name.startswith('List Number'):
            return f"1. {text}"
        else:
            return f"- {text}"

    # Handle quotes
    elif 'Quote' in style_name or 'Block' in style_name:
        return f"> {text}"

    # Apply inline formatting
    md_text = apply_inline_formatting(paragraph)

    return md_text


def apply_inline_formatting(paragraph):
    """
    Apply bold, italic, and other inline formatting to text
    """
    result = []

    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        if run.underline:
            text = f"<u>{text}</u>"

        result.append(text)

    return ''.join(result)


def table_to_markdown(table):
    """
    Convert a table to markdown format (basic version)
    """
    if not table.rows:
        return []

    md_lines = []

    # Header row
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    md_lines.append('| ' + ' | '.join(header_cells) + ' |')

    # Separator
    md_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')

    # Data rows
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        md_lines.append('| ' + ' | '.join(cells) + ' |')

    return md_lines


def extract_docx_metadata(docx_path):
    """
    Extract comprehensive metadata from DOCX file

    Returns:
        Dictionary with document metadata
    """
    docx_path = Path(docx_path)
    doc = docx.Document(docx_path)

    # Core properties
    core_props = doc.core_properties

    metadata = {
        'filename': docx_path.name,
        'file_size_bytes': docx_path.stat().st_size,
        'file_size_mb': round(docx_path.stat().st_size / (1024 * 1024), 2),
        'document_properties': {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'comments': core_props.comments or '',
            'last_modified_by': core_props.last_modified_by or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
            'category': core_props.category or '',
        },
        'structure': {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
        },
        'headings': [],
        'tables_info': [],
        'images': 0,
        'hyperlinks': 0
    }

    # Extract headings
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if 'Heading' in style_name:
            level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
            metadata['headings'].append({
                'level': level,
                'text': para.text.strip(),
                'style': style_name
            })

    # Extract table information
    for i, table in enumerate(doc.tables, 1):
        table_info = {
            'table_number': i,
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'cells': len(table.rows) * (len(table.columns) if table.rows else 0)
        }
        metadata['tables_info'].append(table_info)

    # Count images (embedded shapes)
    try:
        from docx.oxml.ns import qn
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        metadata['images'] = image_count
    except:
        metadata['images'] = 0

    # Count hyperlinks
    hyperlink_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.element.rPr and run.element.rPr.find(qn('w:rStyle')) is not None:
                hyperlink_count += 1
    metadata['hyperlinks'] = hyperlink_count

    return metadata


def extract_headings_from_markdown(markdown_text):
    """
    Extract headings from markdown text
    """
    headings = []
    lines = markdown_text.split('\n')

    for line_num, line in enumerate(lines, 1):
        match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            headings.append({
                'level': level,
                'text': text,
                'line_number': line_num,
                'type': f'H{level}'
            })

    return headings


def analyze_document_structure(markdown_text):
    """
    Analyze document structure from markdown
    """
    lines = markdown_text.split('\n')

    structure = {
        'total_lines': len(lines),
        'non_empty_lines': len([l for l in lines if l.strip()]),
        'headings': len([l for l in lines if l.strip().startswith('#')]),
        'lists': {
            'unordered': len(re.findall(r'^\s*[-*+]\s+', markdown_text, re.MULTILINE)),
            'ordered': len(re.findall(r'^\s*\d+\.\s+', markdown_text, re.MULTILINE))
        },
        'tables': len(re.findall(r'\|[^\n]+\|', markdown_text)),
        'blockquotes': len(re.findall(r'^>\s+', markdown_text, re.MULTILINE)),
        'code_blocks': len(re.findall(r'```', markdown_text)) // 2,
        'bold_text': len(re.findall(r'\*\*([^*]+)\*\*', markdown_text)),
        'italic_text': len(re.findall(r'\*([^*]+)\*', markdown_text)),
    }

    # Word statistics
    words = markdown_text.split()
    structure['word_count'] = len(words)
    structure['char_count'] = len(markdown_text)
    structure['avg_word_length'] = round(sum(len(w) for w in words) / len(words), 2) if words else 0

    return structure


def generate_metadata_report(docx_path, markdown_text, output_path=None):
    """
    Generate comprehensive metadata report for DOCX

    Args:
        docx_path: Path to DOCX file
        markdown_text: Converted markdown text
        output_path: Optional path for metadata JSON file

    Returns:
        Complete metadata dictionary
    """
    print(f"\n{'='*70}")
    print("EXTRACTING METADATA AND DOCUMENT STRUCTURE")
    print(f"{'='*70}\n")

    # Extract all metadata
    print("Extracting DOCX metadata...")
    docx_metadata = extract_docx_metadata(docx_path)

    print("Analyzing headings...")
    headings = extract_headings_from_markdown(markdown_text)

    print("Analyzing document structure...")
    structure = analyze_document_structure(markdown_text)

    # Compile complete report
    report = {
        'extraction_date': datetime.now().isoformat(),
        'source_file': str(docx_path),
        'docx_metadata': docx_metadata,
        'headings': {
            'count': len(headings),
            'by_level': dict(Counter(h['level'] for h in headings)),
            'list': headings
        },
        'structure': structure
    }

    # Save to JSON if output path provided
    if output_path:
        output_path = Path(output_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        print(f"\nMetadata saved to: {output_path}")

    return report


def print_metadata_summary(report):
    """
    Print a human-readable summary of the metadata report
    """
    print(f"\n{'='*70}")
    print("METADATA SUMMARY")
    print(f"{'='*70}\n")

    # Document Info
    docx = report['docx_metadata']
    print(f"Document Information:")
    print(f"  File: {docx['filename']}")
    print(f"  Size: {docx['file_size_mb']} MB")
    print(f"  Paragraphs: {docx['structure']['paragraphs']}")
    print(f"  Tables: {docx['structure']['tables']}")
    print(f"  Sections: {docx['structure']['sections']}")

    # Document Properties
    props = docx['document_properties']
    if any(props.values()):
        print(f"\nDocument Properties:")
        if props['title']:
            print(f"  Title: {props['title']}")
        if props['author']:
            print(f"  Author: {props['author']}")
        if props['subject']:
            print(f"  Subject: {props['subject']}")
        if props['created']:
            print(f"  Created: {props['created']}")
        if props['modified']:
            print(f"  Modified: {props['modified']}")

    # Headings
    headings = report['headings']
    print(f"\nHeadings: {headings['count']} total")
    for level, count in sorted(headings['by_level'].items()):
        print(f"  H{level}: {count}")

    # Structure
    struct = report['structure']
    print(f"\nDocument Structure:")
    print(f"  Lines: {struct['total_lines']:,}")
    print(f"  Words: {struct['word_count']:,}")
    print(f"  Characters: {struct['char_count']:,}")
    print(f"  Lists: {struct['lists']['unordered']} unordered, {struct['lists']['ordered']} ordered")
    print(f"  Tables: {struct['tables']}")
    print(f"  Blockquotes: {struct['blockquotes']}")

    # Media
    print(f"\nMedia:")
    print(f"  Images: {docx['images']}")
    print(f"  Hyperlinks: {docx['hyperlinks']}")

    # Heading Preview
    if headings['list']:
        print(f"\nFirst 5 Headings:")
        for i, h in enumerate(headings['list'][:5], 1):
            indent = "  " * (h['level'] - 1)
            print(f"  {indent}H{h['level']}: {h['text']}")

    print(f"\n{'='*70}")


def batch_convert_docx(docx_directory, output_directory=None, output_format='md'):
    """
    Convert all DOCX files in a directory

    Args:
        docx_directory: Directory containing DOCX files
        output_directory: Directory for output files (optional)
        output_format: 'md' for markdown or 'txt' for plain text
    """
    docx_dir = Path(docx_directory)

    if not docx_dir.exists() or not docx_dir.is_dir():
        raise ValueError(f"Invalid directory: {docx_directory}")

    # Find all DOCX files
    docx_files = list(docx_dir.glob('*.docx'))

    if not docx_files:
        print(f"No DOCX files found in: {docx_directory}")
        return

    print(f"\nFound {len(docx_files)} DOCX files")
    print(f"{'='*70}\n")

    # Setup output directory
    if output_directory:
        output_dir = Path(output_directory)
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = docx_dir

    # Convert each DOCX
    success_count = 0
    failed_files = []

    for i, docx_file in enumerate(docx_files, 1):
        print(f"\n[{i}/{len(docx_files)}] Converting: {docx_file.name}")

        try:
            if output_format == 'md':
                output_file = output_dir / docx_file.with_suffix('.md').name
                convert_docx_to_markdown(docx_file, output_file)
            else:
                output_file = output_dir / docx_file.with_suffix('.txt').name
                convert_docx_to_text(docx_file, output_file)
            success_count += 1
        except Exception as e:
            print(f"  FAILED: {str(e)}")
            failed_files.append(docx_file.name)

    # Summary
    print(f"\n{'='*70}")
    print(f"BATCH CONVERSION COMPLETE")
    print(f"{'='*70}")
    print(f"  Successful: {success_count}/{len(docx_files)}")

    if failed_files:
        print(f"  Failed files:")
        for fname in failed_files:
            print(f"    - {fname}")


def main():
    """
    Command-line interface for DOCX conversion
    """
    parser = argparse.ArgumentParser(
        description='Convert DOCX files to Markdown or Text',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Convert to markdown
  python docx2llm.py document.docx

  # Convert to plain text
  python docx2llm.py document.docx --format txt

  # Convert with custom output name
  python docx2llm.py document.docx -o output.md

  # Extract metadata only
  python docx2llm.py document.docx --metadata-only

  # Batch convert all DOCX files
  python docx2llm.py --batch ./documents/ -o ./output/
        '''
    )

    parser.add_argument(
        'input',
        nargs='?',
        help='Input DOCX file or directory (for batch mode)'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output file or directory (for batch mode)'
    )

    parser.add_argument(
        '--format',
        choices=['md', 'txt'],
        default='md',
        help='Output format: md (markdown) or txt (plain text)'
    )

    parser.add_argument(
        '--batch',
        action='store_true',
        help='Batch convert all DOCX files in input directory'
    )

    parser.add_argument(
        '--metadata',
        action='store_true',
        help='Extract and save metadata'
    )

    parser.add_argument(
        '--metadata-only',
        action='store_true',
        help='Extract only metadata without conversion'
    )

    args = parser.parse_args()

    # Show help if no input provided
    if not args.input:
        parser.print_help()
        sys.exit(1)

    try:
        if args.metadata_only:
            # Extract metadata only
            print("Extracting metadata only...")
            docx_path = Path(args.input)

            # Convert to get markdown for analysis
            markdown_text = convert_docx_to_markdown(docx_path)

            # Generate metadata report
            metadata_output = args.output or docx_path.with_suffix('.metadata.json')
            report = generate_metadata_report(docx_path, markdown_text, metadata_output)
            print_metadata_summary(report)

        elif args.batch:
            # Batch conversion mode
            batch_convert_docx(
                args.input,
                args.output,
                args.format
            )
        else:
            # Single file conversion
            if args.format == 'md':
                markdown_text = convert_docx_to_markdown(args.input, args.output)
            else:
                markdown_text = convert_docx_to_text(args.input, args.output)

            # Extract metadata if requested
            if args.metadata:
                docx_path = Path(args.input)
                metadata_output = docx_path.with_suffix('.metadata.json')
                report = generate_metadata_report(docx_path, markdown_text, metadata_output)
                print_metadata_summary(report)

    except KeyboardInterrupt:
        print("\n\nConversion cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"\nERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
