import logging
from pathlib import Path
from typing import Union, Optional, List, Dict, Any, Tuple
import fitz  # PyMuPDF
from docx import Document
import io
import re
from dataclasses import dataclass, field, asdict
from datetime import datetime
from collections import Counter

# Configure logger
logger = logging.getLogger(__name__)


@dataclass
class PageMetadata:
    """Metadata for a single page."""
    page_number: int
    text: str
    char_count: int
    word_count: int
    urls: List[str] = field(default_factory=list)
    emails: List[str] = field(default_factory=list)
    headings: List[str] = field(default_factory=list)
    headings_structured: List[Dict[str, Any]] = field(default_factory=list)  # Headings with level info
    has_toc: bool = False
    sections: List[str] = field(default_factory=list)  # Section titles if present

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return asdict(self)


@dataclass
class DocumentMetadata:
    """Complete document metadata."""
    file_name: str
    file_type: str
    total_pages: int
    total_chars: int
    total_words: int
    extraction_date: str
    pages: List[PageMetadata] = field(default_factory=list)
    all_urls: List[str] = field(default_factory=list)
    all_emails: List[str] = field(default_factory=list)
    all_headings: List[str] = field(default_factory=list)
    has_toc: bool = False
    toc_page_numbers: List[int] = field(default_factory=list)
    sections: List[str] = field(default_factory=list)  # Document sections
    total_sections: int = 0  # Number of sections found
    has_section_breaks: bool = False  # Whether document has section breaks

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'file_name': self.file_name,
            'file_type': self.file_type,
            'total_pages': self.total_pages,
            'total_chars': self.total_chars,
            'total_words': self.total_words,
            'extraction_date': self.extraction_date,
            'all_urls': self.all_urls,
            'all_emails': self.all_emails,
            'all_headings': self.all_headings,
            'has_toc': self.has_toc,
            'toc_page_numbers': self.toc_page_numbers,
            'sections': self.sections,
            'total_sections': self.total_sections,
            'has_section_breaks': self.has_section_breaks,
            'pages': [page.to_dict() for page in self.pages]
        }


@dataclass
class ExtractionResult:
    """Complete extraction result with text and metadata."""
    text: str
    metadata: DocumentMetadata
    pages: List[PageMetadata]
    
    def get_text_by_page(self, page_number: int) -> Optional[str]:
        """Get text for a specific page."""
        for page in self.pages:
            if page.page_number == page_number:
                return page.text
        return None
    
    def get_pages_range(self, start: int, end: int) -> str:
        """Get combined text from page range."""
        texts = []
        for page in self.pages:
            if start <= page.page_number <= end:
                texts.append(page.text)
        return "\n\n".join(texts)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'text': self.text,
            'metadata': self.metadata.to_dict(),
            'pages': [page.to_dict() for page in self.pages]
        }


class TextPatternDetector:
    """Detect various patterns in text for metadata extraction."""
    
    def __init__(self):
        # URL pattern
        self.url_pattern = re.compile(
            r'(?:https?|ftp)://[^\s]+|www\.[^\s]+',
            re.IGNORECASE
        )
        
        # Email pattern
        self.email_pattern = re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        )
        
        # Heading patterns (all caps, title case, numbered sections)
        # EXPANDED: Added more patterns to catch various heading formats
        self.heading_patterns = [
            # Markdown headings (from enrichment or native markdown)
            re.compile(r'^#{1,6}\s+.+$', re.MULTILINE),

            # ALL CAPS HEADINGS (pure caps, 6+ chars)
            re.compile(r'^[A-Z][A-Z\s]{5,}$', re.MULTILINE),

            # ALL CAPS with # prefix (e.g., "# CONCEPTS IN PRACTICE")
            re.compile(r'^#\s+[A-Z][A-Z\s]{5,}$', re.MULTILINE),

            # Chapter/Section/Part markers
            re.compile(r'^(?:Chapter|Section|Part|Appendix|Article)\s+\d+[:\s]', re.MULTILINE | re.IGNORECASE),

            # Numbered headings: "1. Heading Text"
            re.compile(r'^\d+\.\s+[A-Z][a-zA-Z\s]{3,}$', re.MULTILINE),

            # Numbered headings: "1.1 Heading Text"
            re.compile(r'^\d+\.\d+\s+[A-Z][a-zA-Z\s]{3,}$', re.MULTILINE),

            # Numbered headings: "1.1.1 Heading Text"
            re.compile(r'^\d+\.\d+\.\d+\s+[A-Z][a-zA-Z\s]{3,}$', re.MULTILINE),

            # NEW: Number + symbol + text (e.g., "8 • Strings")
            re.compile(r'^\d+\s+[•·●○▪▫■□\-]\s+[A-Z][a-zA-Z\s]{2,}$', re.MULTILINE),

            # NEW: Table/Figure references (e.g., "Table 8.3", "Figure 2.1")
            re.compile(r'^(?:Table|Figure|Listing|Example)\s+\d+(?:\.\d+)?', re.MULTILINE | re.IGNORECASE),

            # NEW: Underlined headings (heading followed by ===== or -----)
            re.compile(r'^.{5,}\n[=\-]{3,}$', re.MULTILINE),

            # NEW: Bold-style markers (***Heading*** or **Heading**)
            re.compile(r'^\*{2,3}[A-Z][a-zA-Z\s]{3,}\*{2,3}$', re.MULTILINE),
        ]
        
        # TOC patterns
        self.toc_patterns = [
            re.compile(r'(?i)table\s+of\s+contents'),
            re.compile(r'\.{3,}'),  # Dot leaders
            re.compile(r'^.{10,}\.{3,}\s*\d+\s*$', re.MULTILINE),  # Title ... 123
            re.compile(r'^\d+\.\d+\s+.+\s+\d+\s*$', re.MULTILINE),  # 1.2 Title 45
        ]
    
    def extract_urls(self, text: str) -> List[str]:
        """Extract all URLs from text."""
        urls = list(set(self.url_pattern.findall(text)))
        logger.debug(f"Pattern Detection: Found {len(urls)} unique URLs")
        return urls
    
    def extract_emails(self, text: str) -> List[str]:
        """Extract all email addresses from text."""
        emails = list(set(self.email_pattern.findall(text)))
        logger.debug(f"Pattern Detection: Found {len(emails)} unique email addresses")
        return emails
    
    def extract_headings(self, text: str) -> List[str]:
        """Extract potential headings from text."""
        headings = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line or len(line) < 5:
                continue

            # Check against heading patterns
            for pattern in self.heading_patterns:
                if pattern.match(line):
                    headings.append(line)
                    logger.debug(f"Pattern Detection: Detected heading: '{line[:50]}...'")
                    break

        unique_headings = list(set(headings))
        logger.debug(f"Pattern Detection: Found {len(unique_headings)} unique headings")
        return unique_headings
    
    def detect_toc(self, text: str) -> bool:
        """Detect if text contains Table of Contents."""
        for i, pattern in enumerate(self.toc_patterns):
            if pattern.search(text):
                logger.debug(f"Pattern Detection: TOC detected (pattern {i+1}/{len(self.toc_patterns)})")
                return True
        logger.debug("Pattern Detection: No TOC detected")
        return False


def extract_unbiased_pdf_headings(pdf_doc: fitz.Document) -> List[Dict[str, Any]]:
    """
    Extract headings from PDF using font size analysis and visual inference.

    This function uses a two-pass approach:
    1. Pass 1: Collect all text lines with font metadata
    2. Pass 2: Classify headings based on font size relative to body text

    Args:
        pdf_doc: Open PyMuPDF document object

    Returns:
        List of heading dictionaries with keys:
            - page: Page number
            - level: Heading level (H1, H2, H3)
            - text: Heading text
            - font_size: Font size in points
            - bold: Whether text is bold
    """
    logger.debug("Starting unbiased PDF heading extraction")

    lines = []

    # Pass 1: Collect all text lines with metadata
    for page_no, page in enumerate(pdf_doc, start=1):
        page_dict = page.get_text("dict")

        for block in page_dict["blocks"]:
            if block["type"] != 0:  # Skip non-text blocks
                continue

            for line in block["lines"]:
                text = "".join(span["text"] for span in line["spans"]).strip()
                if not text:
                    continue

                sizes = [span["size"] for span in line["spans"]]
                avg_size = sum(sizes) / len(sizes) if sizes else 0

                fonts = {span["font"] for span in line["spans"]}
                bold = any("Bold" in f for f in fonts)

                lines.append({
                    "page": page_no,
                    "text": text,
                    "size": avg_size,
                    "bold": bold
                })

    if not lines:
        logger.debug("No text lines found in PDF")
        return []

    # Determine body font size (most common size)
    size_counts = Counter(round(l["size"], 1) for l in lines)
    body_size = size_counts.most_common(1)[0][0] if size_counts else 12.0
    logger.debug(f"Detected body font size: {body_size}pt")

    headings = []

    # Pass 2: Classify headings based on font size
    for line in lines:
        text = line["text"]
        size = line["size"]
        bold = line["bold"]

        level = None

        # H1 — very large text (1.6x body size)
        if size >= body_size * 1.6:
            level = "H1"

        # H2 — moderately larger (1.3x body size)
        elif size >= body_size * 1.3:
            level = "H2"

        # H3 — slightly larger OR bold isolated line (short text)
        elif size > body_size and bold and len(text.split()) <= 10:
            level = "H3"

        if level:
            headings.append({
                "page": line["page"],
                "level": level,
                "text": text,
                "font_size": round(size, 1),
                "bold": bold
            })

    logger.debug(f"Extracted {len(headings)} headings from PDF using font analysis")
    return headings


def extract_unbiased_docx_headings(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract headings from DOCX using style and font size analysis.

    This function uses a two-pass approach:
    1. Pass 1: Collect paragraph metadata (style, font size, bold)
    2. Pass 2: Classify headings using semantic styles or visual inference

    Args:
        doc: python-docx Document object

    Returns:
        List of heading dictionaries with keys:
            - index: Paragraph index
            - level: Heading level (H1, H2, H3, etc.)
            - text: Heading text
            - font_size: Average font size in points
            - style: Style name
            - bold: Whether text is bold
    """
    logger.debug("Starting unbiased DOCX heading extraction")

    paragraphs = []
    font_sizes = []

    # Pass 1: Collect paragraph metadata
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else ""

        # Calculate average font size from runs
        runs = para.runs
        sizes = [run.font.size.pt for run in runs if run.font.size]
        avg_size = sum(sizes) / len(sizes) if sizes else None

        bold = any(run.bold for run in runs)

        paragraphs.append({
            "index": idx,
            "text": text,
            "style": style,
            "size": avg_size,
            "bold": bold
        })

        if avg_size:
            font_sizes.append(round(avg_size, 1))

    # Determine body font size (most common size)
    body_size = None
    if font_sizes:
        body_size = Counter(font_sizes).most_common(1)[0][0]
        logger.debug(f"Detected body font size: {body_size}pt")

    headings = []

    # Pass 2: Classify headings
    for p in paragraphs:
        text = p["text"]
        style = p["style"]
        size = p["size"]
        bold = p["bold"]

        level = None

        # Primary: Use semantic styles (Heading 1, Heading 2, etc.)
        if style.startswith("Heading"):
            try:
                n = int(style.split()[-1])
                level = f"H{n}"
                logger.debug(f"Found semantic heading: {level} - '{text[:50]}'")
            except ValueError:
                pass

        # Fallback: Visual inference based on font size
        elif body_size and size:
            if size >= body_size * 1.6:
                level = "H1"
            elif size >= body_size * 1.3:
                level = "H2"
            elif size > body_size and bold and len(text.split()) <= 12:
                level = "H3"

        if level:
            headings.append({
                "index": p["index"],
                "level": level,
                "text": text,
                "font_size": size,
                "style": style,
                "bold": bold
            })

    logger.debug(f"Extracted {len(headings)} headings from DOCX using style/font analysis")
    return headings


def extract_text_from_document(
    file_path: Union[str, Path, io.BytesIO],
    file_type: Optional[str] = None,
    encoding: str = 'utf-8',
    extract_metadata: bool = True,
    preserve_page_structure: bool = True
) -> Union[str, ExtractionResult]:
    """
    Extract text content from various document formats with optional metadata.
    
    Args:
        file_path: Path to the document file or BytesIO object
        file_type: Explicit file type ('pdf', 'docx', 'txt'). 
                  If None, inferred from file extension
        encoding: Text encoding for TXT files (default: 'utf-8')
        extract_metadata: If True, return ExtractionResult with metadata
        preserve_page_structure: If True, maintain page-level separation
    
    Returns:
        str or ExtractionResult: Extracted text or complete result with metadata
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If unsupported file type is provided
        Exception: For any processing errors during extraction
        
    Example:
        >>> # Simple text extraction
        >>> text = extract_text_from_document('report.pdf', extract_metadata=False)
        >>> 
        >>> # Full extraction with metadata
        >>> result = extract_text_from_document('report.pdf', extract_metadata=True)
        >>> print(result.metadata.total_pages)
        >>> print(result.pages[0].urls)
    """
    try:
        file_name = "buffer"
        
        # Determine file type
        if isinstance(file_path, (str, Path)):
            file_path = Path(file_path)
            file_name = file_path.name
            
            # Validate file exists
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            # Infer file type from extension if not provided
            if file_type is None:
                file_type = file_path.suffix.lower().lstrip('.')
                logger.debug(f"Inferred file type: {file_type} from {file_path.name}")
            
            logger.info(f"Processing document: {file_path.name} (type: {file_type})")
        else:
            # Handle BytesIO objects
            if file_type is None:
                logger.error("file_type must be specified for BytesIO objects")
                raise ValueError("file_type parameter is required for BytesIO objects")
            logger.info(f"Processing document from buffer (type: {file_type})")
        
        # Process based on file type
        if file_type == 'pdf':
            if extract_metadata:
                result = _extract_from_pdf_with_metadata(file_path, file_name)
            else:
                text = _extract_from_pdf(file_path)
                result = text
            
        elif file_type == 'docx':
            if extract_metadata:
                result = _extract_from_docx_with_metadata(file_path, file_name)
            else:
                text = _extract_from_docx(file_path)
                result = text
            
        elif file_type == 'txt':
            if extract_metadata:
                result = _extract_from_txt_with_metadata(file_path, file_name, encoding)
            else:
                text = _extract_from_txt(file_path, encoding)
                result = text
        else:
            logger.error(f"Unsupported file type: {file_type}")
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: pdf, docx, txt"
            )
        
        # Log extraction statistics
        if isinstance(result, str):
            char_count = len(result)
            word_count = len(result.split())
            logger.info(
                f"Successfully extracted text: {char_count} characters, "
                f"{word_count} words"
            )
        else:
            logger.info(
                f"Successfully extracted document: {result.metadata.total_pages} pages, "
                f"{result.metadata.total_chars} characters, "
                f"{result.metadata.total_words} words"
            )
        
        return result
        
    except FileNotFoundError:
        raise
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from document: {str(e)}") from e


def _enrich_text_with_headings_structured(
    text: str,
    headings: List[str],
    structured_headings: List[Dict[str, Any]]
) -> str:
    """
    Enrich text using structured heading information with accurate level detection.

    This enhanced version uses font size analysis results to apply correct heading levels
    instead of relying on heuristic-based level detection.

    Args:
        text: Original page text
        headings: List of all detected headings (combined from various sources)
        structured_headings: List of structured heading dicts with 'text', 'level', 'font_size', etc.

    Returns:
        Text with headings properly formatted as markdown headers with correct levels
    """
    if not headings:
        return text

    enriched_text = text

    # Create a mapping of heading text to level for quick lookup
    heading_level_map = {}
    for sh in structured_headings:
        heading_level_map[sh['text']] = sh['level']

    # For each heading, ensure it's formatted as a markdown header
    for heading in headings:
        # Check if heading exists in text
        if heading not in enriched_text:
            continue

        # Check if it's already formatted as markdown header
        markdown_variants = [
            f"# {heading}",
            f"## {heading}",
            f"### {heading}",
            f"#### {heading}",
            f"##### {heading}",
            f"###### {heading}"
        ]

        # If already markdown formatted, skip
        is_already_markdown = any(variant in enriched_text for variant in markdown_variants)
        if is_already_markdown:
            continue

        # Get level from structured headings (if available), otherwise use simple detection
        if heading in heading_level_map:
            level_str = heading_level_map[heading]  # e.g., "H1", "H2", "H3"
            level = int(level_str[1])  # Extract numeric level
        else:
            # Fallback to simple detection for pattern-based headings
            level = _detect_heading_level_simple(heading)

        markdown_prefix = "#" * level + " "

        # Escape special regex characters in heading
        escaped_heading = re.escape(heading)

        # Pattern: Heading on its own line (with optional surrounding whitespace)
        pattern = re.compile(
            rf'(^|\n)[ \t]*{escaped_heading}[ \t]*($|\n)',
            re.MULTILINE
        )

        # Replace with markdown formatted version
        replacement = rf'\1{markdown_prefix}{heading}\2'
        enriched_text = pattern.sub(replacement, enriched_text, count=1)

    return enriched_text


def _enrich_text_with_headings(text: str, headings: List[str]) -> str:
    """
    Enrich text by ensuring headings are properly formatted as markdown headers.

    BUG FIX (BUG 1 & BUG 4): This function addresses the issue where headings were extracted
    and stored in metadata but never passed to the downstream pipeline (normalization,
    chunking, embeddings). By converting detected headings to markdown format,
    we ensure they are:
    1. Recognized by the BoundaryDetector in the chunking pipeline (uses ^# pattern)
    2. Preserved during normalization
    3. Visible to RAG systems for better context understanding

    Args:
        text: Original page text
        headings: List of detected headings

    Returns:
        Text with headings properly formatted as markdown headers

    Note: This is the fallback version. Use _enrich_text_with_headings_structured when
    structured heading information is available.
    """
    if not headings:
        return text

    enriched_text = text

    # For each heading, ensure it's formatted as a markdown header if present in text
    for heading in headings:
        # Check if heading exists in text
        if heading not in enriched_text:
            continue

        # Check if it's already formatted as markdown header
        markdown_variants = [
            f"# {heading}",
            f"## {heading}",
            f"### {heading}",
            f"#### {heading}",
            f"##### {heading}",
            f"###### {heading}"
        ]

        # If already markdown formatted, skip
        is_already_markdown = any(variant in enriched_text for variant in markdown_variants)
        if is_already_markdown:
            continue

        # BUG FIX (BUG 4): Use regex for robust heading detection and replacement
        # This ensures headings are on their own line and properly formatted
        # Detect heading level from content
        level = _detect_heading_level_simple(heading)
        markdown_prefix = "#" * level + " "

        # Escape special regex characters in heading
        escaped_heading = re.escape(heading)

        # Pattern 1: Heading on its own line (with optional surrounding whitespace)
        # This matches: \n  Heading  \n or start-of-text Heading \n
        pattern1 = re.compile(
            rf'(^|\n)[ \t]*{escaped_heading}[ \t]*($|\n)',
            re.MULTILINE
        )

        # Replace with markdown formatted version
        # Ensure it's on its own line with proper formatting
        replacement = rf'\1{markdown_prefix}{heading}\2'
        enriched_text = pattern1.sub(replacement, enriched_text, count=1)

    return enriched_text


def _detect_heading_level_simple(heading: str) -> int:
    """
    Detect heading level from heading text.
    Used by _enrich_text_with_headings for simple level detection.

    Returns heading level 1-6 based on content patterns.
    """
    # Chapter/Part = H1
    if re.match(r'^(?:Chapter|Part|CHAPTER|PART)\s+\d+', heading, re.IGNORECASE):
        return 1

    # Section X = H2
    if re.match(r'^(?:Section|SECTION)\s+\d+', heading, re.IGNORECASE):
        return 2

    # Numbered headings (detect level by dots)
    numbered_match = re.match(r'^(\d+(?:\.\d+)*)[.)\s]', heading)
    if numbered_match:
        number = numbered_match.group(1)
        level = len(number.split('.'))
        return min(level, 6)  # Max level is 6

    # ALL CAPS = H1 (if long enough)
    if heading.isupper() and len(heading) > 10:
        return 1

    # Title Case or sentence case = H2 (default for unclear cases)
    return 2


def _extract_from_pdf(file_path: Union[Path, io.BytesIO]) -> str:
    """Extract text from PDF using PyMuPDF (simple mode)."""
    try:
        logger.debug("Starting PDF text extraction")

        if isinstance(file_path, io.BytesIO):
            doc = fitz.open(stream=file_path, filetype="pdf")
        else:
            doc = fitz.open(file_path)

        text_parts = []

        for page_num in range(len(doc)):
            logger.debug(f"Extracting text from page {page_num + 1}/{len(doc)}")
            page = doc[page_num]
            text_parts.append(page.get_text())

        doc.close()

        full_text = "\n\n".join(text_parts)
        logger.debug(f"PDF extraction complete: {len(doc)} pages processed")

        return full_text

    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from PDF: {str(e)}") from e


def _extract_from_pdf_with_metadata(
    file_path: Union[Path, io.BytesIO],
    file_name: str
) -> ExtractionResult:
    """Extract text from PDF with full metadata."""
    try:
        logger.debug("Starting PDF text extraction with metadata")

        if isinstance(file_path, io.BytesIO):
            doc = fitz.open(stream=file_path, filetype="pdf")
        else:
            doc = fitz.open(file_path)

        detector = TextPatternDetector()

        # NEW: Extract structured headings using font size analysis (entire document)
        logger.debug("Extracting structured headings using font analysis")
        structured_headings_all = extract_unbiased_pdf_headings(doc)
        logger.info(f"Font analysis detected {len(structured_headings_all)} headings across document")

        # Group structured headings by page for easy lookup
        structured_headings_by_page = {}
        for heading in structured_headings_all:
            page_num = heading['page']
            if page_num not in structured_headings_by_page:
                structured_headings_by_page[page_num] = []
            structured_headings_by_page[page_num].append(heading)

        pages_metadata = []
        all_text_parts = []
        all_urls = set()
        all_emails = set()
        all_headings = set()
        # BUG FIX: Track headings in order to preserve document hierarchy
        ordered_headings = []  # Preserve order for document-level sections
        toc_pages = []

        for page_num in range(len(doc)):
            logger.debug(f"PDF Extraction: Processing page {page_num + 1}/{len(doc)}")
            page = doc[page_num]
            page_text = page.get_text()

            char_count = len(page_text)
            word_count = len(page_text.split())
            logger.debug(f"PDF Extraction: Page {page_num + 1} - {char_count} chars, {word_count} words")

            # Extract metadata for this page
            urls = detector.extract_urls(page_text)
            emails = detector.extract_emails(page_text)
            headings = detector.extract_headings(page_text)  # Pattern-based headings
            has_toc = detector.detect_toc(page_text)

            # NEW: Get structured headings for this page (from font analysis)
            page_structured_headings = structured_headings_by_page.get(page_num + 1, [])

            # Merge pattern-based and font-based headings (union of both)
            font_based_heading_texts = [h['text'] for h in page_structured_headings]
            combined_headings = list(set(headings + font_based_heading_texts))

            if has_toc:
                toc_pages.append(page_num + 1)
                logger.debug(f"PDF Extraction: Page {page_num + 1} contains Table of Contents")

            if urls:
                logger.debug(f"PDF Extraction: Page {page_num + 1} - Found URLs: {urls[:3]}{'...' if len(urls) > 3 else ''}")
            if emails:
                logger.debug(f"PDF Extraction: Page {page_num + 1} - Found emails: {emails[:3]}{'...' if len(emails) > 3 else ''}")
            if combined_headings:
                logger.debug(f"PDF Extraction: Page {page_num + 1} - Found {len(combined_headings)} headings ({len(headings)} pattern-based, {len(font_based_heading_texts)} font-based)")

            # Update global collections
            all_urls.update(urls)
            all_emails.update(emails)
            all_headings.update(combined_headings)
            # BUG FIX: Preserve heading order for document-level sections
            for heading in combined_headings:
                if heading not in ordered_headings:  # Avoid duplicates while preserving order
                    ordered_headings.append(heading)

            # BUG FIX: Enrich page text with properly formatted headings (with level info)
            # This ensures headings flow through to normalization, chunking, and embeddings
            enriched_page_text = _enrich_text_with_headings_structured(
                page_text,
                combined_headings,
                page_structured_headings
            )

            # Create page metadata
            page_meta = PageMetadata(
                page_number=page_num + 1,
                text=page_text,  # Keep original text in metadata
                char_count=char_count,
                word_count=word_count,
                urls=urls,
                emails=emails,
                headings=combined_headings,  # Combined headings
                headings_structured=page_structured_headings,  # NEW: Structured headings with levels
                has_toc=has_toc,
                sections=[]  # PDF pages don't have section breaks within a page
            )

            pages_metadata.append(page_meta)
            all_text_parts.append(enriched_page_text)  # Use enriched text for downstream processing
        
        doc.close()
        
        # Combine all text
        full_text = "\n\n".join(all_text_parts)
        
        # BUG FIX: Build document-level sections from ordered headings
        # This provides global section structure and continuity across pages
        doc_metadata = DocumentMetadata(
            file_name=file_name,
            file_type='pdf',
            total_pages=len(pages_metadata),
            total_chars=len(full_text),
            total_words=len(full_text.split()),
            extraction_date=datetime.now().isoformat(),
            pages=pages_metadata,
            all_urls=list(all_urls),
            all_emails=list(all_emails),
            all_headings=list(all_headings),
            has_toc=len(toc_pages) > 0,
            toc_page_numbers=toc_pages,
            sections=ordered_headings,  # Document-level sections in order
            total_sections=len(ordered_headings),
            has_section_breaks=len(ordered_headings) > 0  # Has sections if headings exist
        )
        
        logger.debug(
            f"PDF Extraction Complete: {len(pages_metadata)} pages, "
            f"{len(all_urls)} URLs, {len(all_emails)} emails, {len(all_headings)} headings, "
            f"{len(ordered_headings)} sections, "
            f"TOC pages: {toc_pages if toc_pages else 'None'}"
        )
        
        return ExtractionResult(
            text=full_text,
            metadata=doc_metadata,
            pages=pages_metadata
        )
        
    except Exception as e:
        logger.error(f"PDF extraction with metadata failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from PDF: {str(e)}") from e


def _extract_from_docx(file_path: Union[Path, io.BytesIO]) -> str:
    """Extract text from DOCX using python-docx (simple mode)."""
    try:
        logger.debug("Starting DOCX text extraction")

        if isinstance(file_path, io.BytesIO):
            doc = Document(file_path)
        else:
            doc = Document(str(file_path))

        text_parts = []
        current_section_text = []
        section_count = 0

        # Extract text from paragraphs with section awareness
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Check for section break by examining paragraph properties
            # Section breaks in DOCX are indicated by paragraph's section property changes
            if para._element.pPr is not None:
                sect_pr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sect_pr is not None and current_section_text:
                    # Found a section break
                    section_count += 1
                    text_parts.extend(current_section_text)
                    text_parts.append("\n--- SECTION BREAK ---\n")
                    current_section_text = []

            current_section_text.append(para.text)

        # Add remaining section text
        if current_section_text:
            text_parts.extend(current_section_text)

        # Extract text from tables
        if doc.tables:
            text_parts.append("\n--- TABLES ---\n")
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n[Table {table_idx + 1}]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)
        logger.debug(
            f"DOCX extraction complete: {len(doc.paragraphs)} paragraphs, "
            f"{len(doc.tables)} tables, {section_count} section breaks"
        )

        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from DOCX: {str(e)}") from e


def _extract_from_docx_with_metadata(
    file_path: Union[Path, io.BytesIO],
    file_name: str
) -> ExtractionResult:
    """Extract text from DOCX with metadata (treats as single page with section awareness)."""
    try:
        logger.debug("Starting DOCX text extraction with metadata")

        if isinstance(file_path, io.BytesIO):
            doc = Document(file_path)
        else:
            doc = Document(str(file_path))

        detector = TextPatternDetector()

        # NEW: Extract structured headings using style and font size analysis
        logger.debug("Extracting structured headings using style/font analysis")
        structured_headings = extract_unbiased_docx_headings(doc)
        logger.info(f"Style/font analysis detected {len(structured_headings)} headings in DOCX")

        # Track sections and their content
        sections = []
        current_section_text = []
        section_titles = []
        section_count = 0
        has_section_breaks = False

        # Extract text from paragraphs with enhanced metadata
        text_parts = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Check for section break
            if para._element.pPr is not None:
                sect_pr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sect_pr is not None and current_section_text:
                    # Found a section break
                    section_count += 1
                    has_section_breaks = True
                    section_content = "\n".join(current_section_text)
                    sections.append(section_content)
                    text_parts.extend(current_section_text)
                    text_parts.append("\n--- SECTION BREAK ---\n")
                    current_section_text = []

            # Detect if this paragraph is a heading
            if para.style and para.style.name.startswith('Heading'):
                section_titles.append(para.text.strip())

            current_section_text.append(para.text)

        # Add remaining section text
        if current_section_text:
            section_content = "\n".join(current_section_text)
            sections.append(section_content)
            text_parts.extend(current_section_text)

        # Extract text from tables with better formatting
        table_texts = []
        if doc.tables:
            text_parts.append("\n--- TABLES ---\n")
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n[Table {table_idx + 1}]")
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        formatted_row = " | ".join(row_text)
                        text_parts.append(formatted_row)
                        table_content.append(formatted_row)
                table_texts.append("\n".join(table_content))

        full_text = "\n".join(text_parts)

        # Extract metadata
        urls = detector.extract_urls(full_text)
        emails = detector.extract_emails(full_text)
        headings = detector.extract_headings(full_text)  # Pattern-based headings

        # NEW: Merge all heading sources (styles, patterns, structured font-based)
        # BUG FIX: Preserve order when combining headings from multiple sources
        all_headings_ordered = []
        seen_headings = set()

        # First add section_titles (from DOCX styles) in order
        for heading in section_titles:
            if heading not in seen_headings:
                all_headings_ordered.append(heading)
                seen_headings.add(heading)

        # Then add structured headings (from font/style analysis)
        font_based_heading_texts = [h['text'] for h in structured_headings]
        for heading in font_based_heading_texts:
            if heading not in seen_headings:
                all_headings_ordered.append(heading)
                seen_headings.add(heading)

        # Finally add pattern-detected headings
        for heading in headings:
            if heading not in seen_headings:
                all_headings_ordered.append(heading)
                seen_headings.add(heading)

        all_headings = all_headings_ordered
        logger.debug(f"Combined headings: {len(section_titles)} style-based, {len(font_based_heading_texts)} font-based, {len(headings)} pattern-based = {len(all_headings)} total")

        # BUG FIX: Enrich text with properly formatted headings using structured info
        # This ensures headings flow through to normalization, chunking, and embeddings
        full_text = _enrich_text_with_headings_structured(full_text, all_headings, structured_headings)

        has_toc = detector.detect_toc(full_text)

        # BUG FIX: Use all_headings (not just section_titles) for complete section hierarchy
        # Create single page metadata with section information
        page_meta = PageMetadata(
            page_number=1,
            text=full_text,
            char_count=len(full_text),
            word_count=len(full_text.split()),
            urls=urls,
            emails=emails,
            headings=all_headings,  # Combined headings
            headings_structured=structured_headings,  # NEW: Structured headings with levels
            has_toc=has_toc,
            sections=all_headings  # Use all headings, not just section_titles
        )

        # BUG FIX: Build document-level sections from all headings (styles + detected)
        # This provides complete global section structure
        doc_metadata = DocumentMetadata(
            file_name=file_name,
            file_type='docx',
            total_pages=1,
            total_chars=len(full_text),
            total_words=len(full_text.split()),
            extraction_date=datetime.now().isoformat(),
            pages=[page_meta],
            all_urls=urls,
            all_emails=emails,
            all_headings=all_headings,
            has_toc=has_toc,
            toc_page_numbers=[1] if has_toc else [],
            sections=all_headings,  # Use all headings for complete hierarchy
            total_sections=len(all_headings),
            has_section_breaks=has_section_breaks or len(all_headings) > 0
        )

        logger.debug(
            f"DOCX extraction complete: {len(doc.paragraphs)} paragraphs, "
            f"{len(doc.tables)} tables, {len(urls)} URLs, "
            f"{len(all_headings)} headings, {doc_metadata.total_sections} sections, "
            f"{'with' if has_section_breaks else 'without'} section breaks"
        )

        return ExtractionResult(
            text=full_text,
            metadata=doc_metadata,
            pages=[page_meta]
        )

    except Exception as e:
        logger.error(f"DOCX extraction with metadata failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from DOCX: {str(e)}") from e


def _extract_from_txt(file_path: Union[Path, io.BytesIO], encoding: str) -> str:
    """Extract text from TXT file using native Python (simple mode)."""
    try:
        logger.debug(f"Starting TXT text extraction with encoding: {encoding}")
        
        if isinstance(file_path, io.BytesIO):
            text = file_path.read().decode(encoding)
        else:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
        
        logger.debug("TXT extraction complete")
        return text
        
    except UnicodeDecodeError as e:
        logger.error(
            f"Encoding error with {encoding}: {str(e)}. "
            f"Try a different encoding."
        )
        raise Exception(
            f"Failed to decode text file with encoding '{encoding}': {str(e)}"
        ) from e
    except Exception as e:
        logger.error(f"TXT extraction failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from TXT: {str(e)}") from e


def _extract_from_txt_with_metadata(
    file_path: Union[Path, io.BytesIO],
    file_name: str,
    encoding: str
) -> ExtractionResult:
    """Extract text from TXT with metadata (treats as single page)."""
    try:
        logger.debug(f"Starting TXT text extraction with metadata, encoding: {encoding}")
        
        if isinstance(file_path, io.BytesIO):
            text = file_path.read().decode(encoding)
        else:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
        
        detector = TextPatternDetector()

        # Extract metadata
        urls = detector.extract_urls(text)
        emails = detector.extract_emails(text)
        headings = detector.extract_headings(text)

        # BUG FIX: Enrich text with properly formatted headings
        # This ensures headings flow through to normalization, chunking, and embeddings
        text = _enrich_text_with_headings(text, headings)

        has_toc = detector.detect_toc(text)
        
        # BUG FIX: Use headings as sections for document-level hierarchy
        # Create single page metadata
        page_meta = PageMetadata(
            page_number=1,
            text=text,
            char_count=len(text),
            word_count=len(text.split()),
            urls=urls,
            emails=emails,
            headings=headings,
            has_toc=has_toc,
            sections=headings  # Use detected headings as sections
        )

        # BUG FIX: Build document-level sections from headings
        # This provides global section structure
        doc_metadata = DocumentMetadata(
            file_name=file_name,
            file_type='txt',
            total_pages=1,
            total_chars=len(text),
            total_words=len(text.split()),
            extraction_date=datetime.now().isoformat(),
            pages=[page_meta],
            all_urls=urls,
            all_emails=emails,
            all_headings=headings,
            has_toc=has_toc,
            toc_page_numbers=[1] if has_toc else [],
            sections=headings,  # Document-level sections from headings
            total_sections=len(headings),
            has_section_breaks=len(headings) > 0  # Has sections if headings exist
        )
        
        logger.debug(
            f"TXT extraction complete: {len(urls)} URLs, {len(headings)} headings, "
            f"{len(headings)} sections"
        )
        
        return ExtractionResult(
            text=text,
            metadata=doc_metadata,
            pages=[page_meta]
        )
        
    except Exception as e:
        logger.error(f"TXT extraction with metadata failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract text from TXT: {str(e)}") from e


# Example usage
if __name__ == "__main__":
    # Configure logging for demonstration
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    print("=" * 70)
    print("EXAMPLE 1: Simple text extraction")
    print("=" * 70)
    try:
        text = extract_text_from_document('example.pdf', extract_metadata=False)
        print(f"Extracted {len(text)} characters")
        print(text[:500] + "...")
    except Exception as e:
        print(f"Error: {e}")
    
    print("\n" + "=" * 70)
    print("EXAMPLE 2: Full extraction with metadata (PDF)")
    print("=" * 70)
    try:
        result = extract_text_from_document('example.pdf', extract_metadata=True)

        print(f"\nDocument: {result.metadata.file_name}")
        print(f"Type: {result.metadata.file_type}")
        print(f"Total Pages: {result.metadata.total_pages}")
        print(f"Total Characters: {result.metadata.total_chars}")
        print(f"Total Words: {result.metadata.total_words}")
        print(f"URLs Found: {len(result.metadata.all_urls)}")
        print(f"Emails Found: {len(result.metadata.all_emails)}")
        print(f"Headings Found: {len(result.metadata.all_headings)}")
        print(f"Has TOC: {result.metadata.has_toc}")

        if result.metadata.all_urls:
            print(f"\nURLs: {result.metadata.all_urls[:3]}")

        if result.metadata.all_headings:
            print(f"\nHeadings: {result.metadata.all_headings[:5]}")

        # Access specific page
        print(f"\nPage 1 has {result.pages[0].word_count} words")
        print(f"Page 1 URLs: {result.pages[0].urls}")

        # Get text by page range
        if result.metadata.total_pages >= 2:
            page_range_text = result.get_pages_range(1, 2)
            print(f"\nPages 1-2 combined: {len(page_range_text)} characters")

    except Exception as e:
        print(f"Error: {e}")

    print("\n" + "=" * 70)
    print("EXAMPLE 3: DOCX extraction with section awareness")
    print("=" * 70)
    try:
        result = extract_text_from_document('example.docx', extract_metadata=True)

        print(f"\nDocument: {result.metadata.file_name}")
        print(f"Type: {result.metadata.file_type}")
        print(f"Total Sections: {result.metadata.total_sections}")
        print(f"Has Section Breaks: {result.metadata.has_section_breaks}")
        print(f"Total Characters: {result.metadata.total_chars}")
        print(f"Total Words: {result.metadata.total_words}")
        print(f"URLs Found: {len(result.metadata.all_urls)}")
        print(f"Headings Found: {len(result.metadata.all_headings)}")

        if result.metadata.sections:
            print(f"\nSection Titles Found:")
            for idx, section_title in enumerate(result.metadata.sections[:5], 1):
                print(f"  {idx}. {section_title}")

        if result.metadata.all_headings:
            print(f"\nAll Headings: {result.metadata.all_headings[:5]}")

    except Exception as e:
        print(f"Error: {e}")